"""Build and export the elemental / peak-fit tables used in manuscripts."""

from __future__ import annotations

import io
from typing import Sequence

import pandas as pd

from .parser import Dataset, PeakRow

# metric key -> (column label, number of decimals)
METRICS: dict[str, tuple[str, int]] = {
    "peak_be": ("Peak BE (eV)", 2),
    "fwhm": ("FWHM (eV)", 2),
    "height": ("Height (CPS)", 0),
    "area_p": ("Area (CPS·eV)", 0),
    "area_n": ("Area (N) TPP-2M", 2),
    "weight_pct": ("Weight (%)", 2),
    "atomic_pct": ("Atomic (%)", 2),
}

DEFAULT_METRICS = ["peak_be", "weight_pct"]


def _fmt(value: float | None, decimals: int) -> str:
    if value is None or pd.isna(value):
        return "—"
    return f"{value:.{decimals}f}"


# which Peak Table rows to prefer when a sample was loaded from several files
SOURCE_CHOICES = {
    "core": "Core-level taramaları (önerilen)",
    "survey": "Survey taraması",
    "fit": "Dekonvolüsyon bileşenleri",
    "any": "Farketmez (en şiddetli)",
}


def element_order(datasets: Sequence[Dataset], source: str = "any") -> list[str]:
    """Elements in first-seen order across all loaded samples.

    Falls back to every source if the requested one is absent, so selecting
    "core" on a survey-only project still lists the elements.
    """
    def collect(wanted: str) -> list[str]:
        order: list[str] = []
        for ds in datasets:
            for peak in ds.peaks:
                if wanted not in ("any", peak.source):
                    continue
                if peak.element not in order:
                    order.append(peak.element)
        return order

    return collect(source) or collect("any")


def _pick_peak(peaks: list[PeakRow], element: str, mode: str,
               source: str = "any") -> PeakRow | None:
    """Reduce several rows of one element to the single one the table needs."""
    candidates = [p for p in peaks if p.element.lower() == element.lower()]
    if source != "any":
        preferred = [p for p in candidates if p.source == source]
        candidates = preferred or candidates
    if not candidates:
        return None
    if mode == "main" or len(candidates) == 1:
        return max(candidates, key=lambda p: (p.height or p.area_p or 0))
    return candidates[0]


def summary_table(
    datasets: Sequence[Dataset],
    elements: Sequence[str] | None = None,
    metrics: Sequence[str] = tuple(DEFAULT_METRICS),
    labels: Sequence[str] | None = None,
    peak_choice: str = "main",
    source: str = "core",
) -> pd.DataFrame:
    """Table 4 layout: samples down the side, elements across the top.

    Each sample contributes one row per selected metric.
    """
    elements = list(elements or element_order(datasets, source))
    labels = list(labels or [ds.label for ds in datasets])

    records = []
    for ds, label in zip(datasets, labels):
        for metric in metrics:
            metric_label, decimals = METRICS[metric]
            row = {"Sample": label, "Parameter": metric_label}
            for element in elements:
                peak = _pick_peak(ds.peaks, element, peak_choice, source)
                row[element] = _fmt(getattr(peak, metric, None) if peak else None, decimals)
            records.append(row)

    df = pd.DataFrame.from_records(records)
    if df.empty:
        return df
    return df.set_index(["Sample", "Parameter"])


def detailed_table(
    datasets: Sequence[Dataset],
    metrics: Sequence[str] = ("peak_be", "fwhm", "area_p", "weight_pct", "atomic_pct"),
    labels: Sequence[str] | None = None,
    elements: Sequence[str] | None = None,
    source: str = "fit",
) -> pd.DataFrame:
    """One row per fitted component - the usual deconvolution table."""
    labels = list(labels or [ds.label for ds in datasets])
    wanted = {e.lower() for e in elements} if elements else None

    records = []
    for ds, label in zip(datasets, labels):
        rows_for_sample = [p for p in ds.peaks if source in ("any", p.source)] or ds.peaks
        for peak in rows_for_sample:
            if wanted and peak.element.lower() not in wanted:
                continue
            row = {
                "Sample": label,
                "Region": peak.element,
                "Component": peak.component or "—",
            }
            for metric in metrics:
                metric_label, decimals = METRICS[metric]
                row[metric_label] = _fmt(getattr(peak, metric, None), decimals)
            records.append(row)
    return pd.DataFrame.from_records(records)


def chemical_state_table(datasets: Sequence[Dataset],
                         labels: Sequence[str] | None = None) -> pd.DataFrame:
    """Avantage's chemical-state assignments, one row per assignment."""
    labels = list(labels or [ds.label for ds in datasets])
    records = []
    for ds, label in zip(datasets, labels):
        for state in ds.chemical_states:
            records.append({
                "Sample": label,
                "Region": state.region,
                "Assignment": state.detail,
                "Chemical state": state.compound,
                "Confidence": "high" if state.confident else "low",
            })
    return pd.DataFrame.from_records(records)


# ---------------------------------------------------------------------------
# flattening for export
# ---------------------------------------------------------------------------
def flatten(df: pd.DataFrame, merge_sample_column: bool = True) -> tuple[list[str], list[list[str]]]:
    """Return (header, rows) with repeated sample names blanked out."""
    if df.empty:
        return [], []
    work = df.reset_index() if isinstance(df.index, pd.MultiIndex) else df.copy()
    header = [str(c) for c in work.columns]
    rows = [[("" if pd.isna(v) else str(v)) for v in rec] for rec in work.itertuples(index=False)]

    if merge_sample_column and header and header[0] == "Sample":
        previous = None
        for row in rows:
            if row[0] == previous:
                row[0] = ""
            else:
                previous = row[0]
    return header, rows


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------
def to_docx(df: pd.DataFrame, caption: str = "", font: str = "Times New Roman",
            font_size: float = 9.0, merge_sample_column: bool = True) -> bytes:
    """Three-line (booktabs) journal table: no vertical rules, rules top/mid/bottom."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    header, rows = flatten(df, merge_sample_column)
    if not header:
        raise ValueError("Tablo boş.")

    doc = Document()
    if caption:
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.bold = True
        run.font.name = font
        run.font.size = Pt(font_size)

    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.autofit = True

    def set_border(cell, edge: str, size: int = 8) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), "000000")

    def write(cell, text: str, bold: bool = False) -> None:
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.bold = bold
        run.font.name = font
        run.font.size = Pt(font_size)

    for j, name in enumerate(header):
        cell = table.rows[0].cells[j]
        write(cell, name, bold=True)
        set_border(cell, "top", 12)
        set_border(cell, "bottom", 8)

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            write(table.rows[i].cells[j], value)
    for cell in table.rows[-1].cells:
        set_border(cell, "bottom", 12)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Excel export
# ---------------------------------------------------------------------------
def to_xlsx(df: pd.DataFrame, sheet_name: str = "XPS Table", caption: str = "",
            merge_sample_column: bool = True) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, Side

    header, rows = flatten(df, merge_sample_column)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "Table"

    start = 1
    if caption:
        ws.cell(row=1, column=1, value=caption).font = Font(bold=True, size=11)
        start = 3

    thin = Side(style="thin", color="000000")
    medium = Side(style="medium", color="000000")

    for j, name in enumerate(header, start=1):
        cell = ws.cell(row=start, column=j, value=name)
        cell.font = Font(bold=True)
        cell.border = Border(top=medium, bottom=thin)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for i, row in enumerate(rows, start=start + 1):
        for j, value in enumerate(row, start=1):
            try:
                cell_value: object = float(value)
            except (TypeError, ValueError):
                cell_value = value
            cell = ws.cell(row=i, column=j, value=cell_value)
            cell.alignment = Alignment(horizontal="center" if j > 2 else "left")

    for j in range(1, len(header) + 1):
        ws.cell(row=start + len(rows), column=j).border = Border(bottom=medium)

    for j, name in enumerate(header, start=1):
        longest = max([len(str(name))] + [len(r[j - 1]) for r in rows]) if rows else len(str(name))
        ws.column_dimensions[ws.cell(row=start, column=j).column_letter].width = min(28, longest + 4)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def to_latex(df: pd.DataFrame, caption: str = "", label: str = "tab:xps") -> str:
    header, rows = flatten(df)
    if not header:
        return ""
    align = "l" * 2 + "c" * (len(header) - 2) if len(header) > 2 else "l" * len(header)
    lines = [
        r"\begin{table}[htbp]",
        r"\centering",
        rf"\caption{{{caption}}}" if caption else r"\caption{}",
        rf"\label{{{label}}}",
        rf"\begin{{tabular}}{{{align}}}",
        r"\hline",
        " & ".join(header) + r" \\",
        r"\hline",
    ]
    lines += [" & ".join(r) + r" \\" for r in rows]
    lines += [r"\hline", r"\end{tabular}", r"\end{table}"]
    return "\n".join(lines)
